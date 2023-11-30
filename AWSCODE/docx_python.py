import docx
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import subprocess
import fitz  # PyMuPDF

# Create a PDF with only the text you want to count
def create_counting_doc(docx_file, text_to_count):
    #initialize Doc object
    doc = docx.Document()
    p = doc.add_paragraph()
    
    #Left Justify Paragraph
    p_format = p.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_format.line_spacing = 1
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)

    #Create run for styling
    r = p.add_run(text_to_count)
    font = r.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.bold = True
    
    #Save
    doc.save(docx_file)

# Convert the DOCX to a PDF
def convert_to_pdf(docx_file):
    # Command to run soffice for conversion
    command = [
        "soffice", 
        "--headless", 
        "--convert-to", "pdf", 
        docx_file,
        "--outdir", "/mountedvolume/"
    ]

    # Execute the command
    subprocess.run(command, check=True)

def count_lines_in_pdf(pdf_file):
    with fitz.open(pdf_file) as doc:
        text = ""
        for page in doc:
            text += page.get_text()
        text = text.rstrip('\n')  # Strip trailing newlines
        lines = [line for line in text.split('\n') if line.strip()]  # Ignore empty lines
        return len(lines)


# Get and print the current working directory
current_directory = os.getcwd()
print("Current working directory:", current_directory)

file_name = "/mountedvolume/TestDoc.docx"

''' 
#Create DOCX file, paragraph, and run
doc = docx.Document()
p = doc.add_paragraph()

#Left Justify Paragraph
p_format = p.paragraph_format
p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_format.line_spacing = 1

#Font Size, Bold, Name
r = p.add_run('Petty Officer GLIESMANN is cited for superior performance of duty while serving in multiple positions at Sector Baltimore, Maryland from April 2011 to June 2015. Reasons_Start While assigned to the Port State Control Division from April 2011 to April 2013, he displayed distinguished performance conducting over 70 examinations resulting in four vessel detentions. Reasons_End Petty Officer GLIESMANN’s diligence, perseverance, and devotion to duty are most heartily commended and are in keeping with the highest traditions of the United States Coast Guard. Petty Officer GLIESMANN is cited for superior performance of duty while serving in multiple positions at Sector Baltimore, Maryland from April 2011 to June 2015. Reasons_Start While assigned to the Port State Control Division from April 2011 to April 2013, he displayed distinguished performance conducting over 70 examinations resulting in four vessel detentions. Reasons_End Petty Officer GLIESMANN’s diligence, perseverance, and devotion to duty are most heartily commended and are in keeping with the highest traditions of the United States Coast Guard.')
font = r.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.bold = True

#Line and paragraph Spacing
p_format.line_spacing = 1
p_format.space_after = Pt(0)
p_format.space_before = Pt(0)

# Add blank lines (paragraph break)
doc.add_paragraph()  
doc.add_paragraph()  

#add p2 
p2 = doc.add_paragraph()

#Font Size, Bold, Name
r2 = p2.add_run('Lorem ipsum dolor sit amet.Lorem t amet amet.')
font2 = r2.font
font2.name = 'Times New Roman'
font2.size = Pt(12)
font2.bold = True

#Left Justify Paragraph
p_format2 = p2.paragraph_format
p_format2.alignment = WD_ALIGN_PARAGRAPH.LEFT

#Line and paragraph Spacing
p_format2.line_spacing = 1
p_format2.space_after = Pt(0)
p_format2.space_before = Pt(0)

doc.save(file_name)
'''

# Text you want to count
text_to_count = 'Petty Officer GLIESMANN is cited for superior performance of duty...'

# File paths
file_name_to_count = "/mountedvolume/CountingDoc.docx"
file_name_full = "/mountedvolume/FullDoc.docx"


# Create document for line counting
create_counting_doc(file_name_to_count, text_to_count)


convert_to_pdf(file_name_to_count)
# Assuming the output file name is the same but with .pdf
pdf_file_to_count = file_name_to_count.rsplit('.', 1)[0] + '.pdf'


# Count lines in the PDF
line_count = count_lines_in_pdf(pdf_file_to_count)
print(f"Number of lines: {line_count}")

#Ill probably end up needing GPT to send a unique ID param so that everytime a user queries they dont create a new .docx