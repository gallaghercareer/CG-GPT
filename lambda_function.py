import docx
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import subprocess
import json
import fitz  # PyMuPD

# Create a PDF with only the text you want to count
def create_counting_doc(docx_file, text_to_count):
    #initialize Doc object
    doc = docx.Document()
    p = doc.add_paragraph()
    
    #Left Justify Paragraph
    p_format = p.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_format.line_spacing = 1
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)

    #Create run for styling
    r = p.add_run(text_to_count)
    font = r.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.bold = True
    
    #Save
    doc.save(docx_file)

# Convert the DOCX to a PDF
def convert_to_pdf(docx_file):
       # Set HOME environment variable to /tmp
    os.environ['HOME'] = '/tmp'

    # Command to run soffice for conversion
    command = [
        "soffice", 
        "--headless", 
        "--convert-to", "pdf", 
        docx_file,
        "--outdir", "/tmp/"
    ]

    # Try executing the command, with a retry on failure
    max_retries = 3
    for attempt in range(max_retries):
        try:
            subprocess.run(command, check=True)
            break  # If successful, break out of the loop
        except subprocess.CalledProcessError as e:
            if attempt < max_retries - 1:  # Retry if not the last attempt
                continue
            else:
                raise  # Re-raise the exception if out of retries

def count_lines_in_pdf(pdf_file):
    with fitz.open(pdf_file) as doc:
        text = ""
        for page in doc:
            text += page.get_text()
        text = text.rstrip('\n')  # Strip trailing newlines
        lines = [line for line in text.split('\n') if line.strip()]  # Ignore empty lines
        return len(lines)
    
def handler(event, context):
    try:

        body = json.loads(event['body'])
        text_to_count = body['paragraph']

        # Create a temporary DOCX file with the provided text
        temp_docx_file = "/tmp/CountingDoc.docx"
        create_counting_doc(temp_docx_file, text_to_count)
        
        # Convert the temporary DOCX to PDF
        convert_to_pdf(temp_docx_file)
        temp_pdf_file = temp_docx_file.rsplit('.', 1)[0] + '.pdf'
        
        # Count lines in the PDF
        line_count = count_lines_in_pdf(temp_pdf_file)
        

        # Convert the line count to a string and embed in a JSON structure
        response_body = {
            "line_count": line_count,           
        }
            # Return the line count as the response
        return {
            "isBase64Encoded": False,
            "statusCode": 200,
            "headers": { "Content-Type": "application/json" },
            "body": json.dumps(response_body)           
        }
        
    except Exception as e:
        # Correctly format the error response according to Lambda proxy integration requirements
        error_response = {
            "errorMessage": str(e)
        }
        return {
            'statusCode': 500,
            'body': json.dumps(error_response),
            "isBase64Encoded": False,
            "headers": { "Content-Type": "application/json" }
        }